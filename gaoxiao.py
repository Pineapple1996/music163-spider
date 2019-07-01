# -*- coding: utf-8 -*-

# -------------------------------------------------
# @Time    : 7/1/19 9:24 AM
# @Author  : 蒋默然
# @File    : gaoxiao.py
# -------------------------------------------------
import os
import time
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import requests
import xlrd
from pyquery import PyQuery as pq


def do_search(keyword):
    url = 'https://gaokao.chsi.com.cn/zsgs/zhangcheng/listVerifedZszc.do'
    data = {
        'method': 'listInfoByYxmc',
        'yxmc': keyword
    }
    headers = {
        'Pragma': 'no-cache',
        'Referer': 'https://gaokao.chsi.com.cn/zsgs/zhangcheng/listVerifedZszc--method-index,lb-1.dhtml',
        'Upgrade-Insecure-Requests': '1',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3',
        'Content-Type': 'application/x-www-form-urlencoded',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/75.0.3770.80 Safari/537.36'
    }
    for i in range(10):
        try:
            resp = requests.post(url, data=data, headers=headers, timeout=10)
            pq_html = pq(resp.text)
            break
        except Exception as e:
            print(e)

    if not pq_html('.zszcdel .ablue').attr('href'):
        print('extract error ' + keyword)
        with open('error.txt','a') as f:
            f.write(keyword)
            f.write('\n')
        return None
    detail_url = 'https://gaokao.chsi.com.cn' + pq_html('.zszcdel .ablue').attr('href')
    print(keyword, detail_url)
    return detail_url


def get_detail(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/75.0.3770.80 Safari/537.36'
    }
    for i in range(10):
        try:
            resp = requests.get(url, headers=headers, timeout=5)
            break
        except Exception as e:
            print(e)
    pq_html = pq(resp.text)
    content = pq_html('.content').text()
    title = pq_html('h2.center').text()
    return title, content


def write_doc(title, content, key):
    # with open('text.docx','w') as f:
    #     f.write(content)
    document = Document()
    style = document.styles['Normal']
    # 设置西文字体
    style.font.name = 'Times New Roman'
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_ = document.add_heading(title, 1)  # 插入标题
    title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph(content)  # 插入段落
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    document.save('{}-{}-2019年本科招生简章.docx'.format(key['id'], key['name']))  # 保存文档


def get_excel_key(path):
    workbook = xlrd.open_workbook(path)
    sheet1 = workbook.sheet_by_name('sheet1')
    l = []
    for i in range(3, sheet1.nrows):
        row = sheet1.row_values(i)
        print(row)
        if not isinstance(row[0], float):
            print('------')
            continue
        de = {'name': row[1], 'id': str(int(row[0])).zfill(4)}
        l.append(de)
    # print(l)
    return l


if __name__ == '__main__':
    path = r'H:\project\music163-spider\高校名单.xls'
    key_list = get_excel_key(path)

    for i in key_list:
        key = i['name']
        url = do_search(key)
        if not url:
            continue
        time.sleep(1)
        title, content = get_detail(url)
        write_doc(title, content, i)
